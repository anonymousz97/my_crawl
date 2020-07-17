from bs4 import BeautifulSoup
import requests
import pandas as pd
import re
from lxml.html import fromstring
from docx import Document
from docx.shared import Pt
import shutil
from docx.shared import Inches


from docx.enum.dml import MSO_THEME_COLOR_INDEX

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink


def exist_img(width):
    return width is not None

def exist_table(width):
    return width is None
#init doc
def get(link_de_thi):
    document = Document()

    #doc style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Open Sans'
    font.size = Pt(11.5)
    sections = document.sections
    margin = 0.5
    for section in sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
    #add link
    p = document.add_paragraph('Link : ')
    p.add_run(link_de_thi).bold=True

    #request link
    r = requests.get(link_de_thi)
    soup = BeautifulSoup(r.content, 'html.parser')

    list_img = []
    for temp in soup.find_all('img',width=exist_img):
        list_img.append(temp['src'])

    list_question_img = []
    for temp in list_img:
        t = r.text[r.text[:r.text.find(temp)].rfind('Câu'):r.text[:r.text.find(temp)].rfind('Câu')+6]
        url = temp
        response = requests.get(url, stream=True)
        with open(t+'.png', 'wb') as out_file:
            shutil.copyfileobj(response.raw, out_file)
        del response
        list_question_img.append(t)

    list_table = []
    for temp in soup.find_all('table',style=exist_table):
        list_table.append(str(temp))

    list_question_table = []
    for temp in list_table:
        t = r.text[r.text[:r.text.find(temp)].rfind('Câu'):r.text[:r.text.find(temp)].rfind('Câu')+6]
        list_question_table.append(t)


    spans = soup.find_all('p')
    i = 0
    list_question_solve = []
    for span in spans[3:-8]:
        if "Xem giải" in span.text:
            list_question_solve.append(span.text[11:17])
        if span.text.startswith('A.')==True or span.text.startswith('B.')==True or span.text.startswith('C.')==True or span.text.startswith('D.')==True or span.text.startswith('Câu')==True or span.text.startswith('(Xem giải)')==True:
            result=""
            if span.text.startswith('A.')==True or span.text.startswith('B.')==True or span.text.startswith('C.')==True or span.text.startswith('D.')==True or span.text.startswith('Câu')==True:
                result = span.text
            elif span.text.startswith('(Xem giải)')==True:
                result = span.text[11:]
            if span.text.startswith('A.')==True or span.text.startswith('B.')==True or span.text.startswith('C.')==True or span.text.startswith('D.')==True:
                p = document.add_paragraph(result)
            else:
                p = document.add_paragraph('')
                p.add_run(result[:6]).bold=True
                p.add_run(result[6:])
                if result[:6] in list_question_img:
                    document.add_picture(result[:6]+'.png')
                if result[:6] in list_question_table:
                    df = pd.read_html(str(temp),header=0)
                    #print(df)
                    t = document.add_table(df[0].shape[0]+1, df[0].shape[1])
                    t.style = 'TableGrid'
                    # add the header rows.
                    for j in range(df[0].shape[-1]):
                        t.cell(0,j).text = df[0].columns[j]
                    # add the rest of the data frame
                    for i in range(df[0].shape[0]):
                        for j in range(df[0].shape[-1]):
                            t.cell(i+1,j).text = str(df[0].values[i,j])
        else:
            if span.text.startswith("⇒") or span.text.startswith("[…]"):
                continue
            p = document.add_paragraph(span.text)

    document.save('de.docx')

    # save solve

    document2 = Document()

    #doc style
    style = document2.styles['Normal']
    font = style.font
    font.name = 'Open Sans'
    font.size = Pt(11.5)
    sections = document2.sections
    margin = 0.5
    for section in sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)
    # for i in list_question_solve:
    #     print(i)
    # print(len(list_question_solve))
                    
    list_link = [tag['href'] for tag in soup.select('p a[href]')]
    table = soup.find('table')
    if table != None:
        df = pd.read_html(str(table),header=0)
        t = document2.add_table(df[0].shape[0]+1, df[0].shape[1])
        t.style = 'TableGrid'
        # add the header rows.
        for j in range(df[0].shape[-1]):
            t.cell(0,j).text = df[0].columns[j]
        # add the rest of the data frame
        for i in range(df[0].shape[0]):
            for j in range(df[0].shape[-1]):
                t.cell(i+1,j).text = str(df[0].values[i,j])
    #de = soup.findAll('div',attrs={"class":"entry-content clearfix"})

    i=0
    for link in list_link:
        if "http://hoctap.dvtienich.com/questions/question/" in link:
            i+=1
            r = requests.get(link)
            soup = BeautifulSoup(r.content,'html.parser')
            table = soup.findAll('div',attrs={"class":"question-content ap-q-content"})
            for x in table:
                p2 = document2.add_paragraph('')
                p2.add_run(list_question_solve[i-1]+" ").bold = True
                temp = x.findAll('p')
                for y in temp:
                    t = y.text
                    p2.add_run(t+"\n")
            table = soup.findAll('div',attrs={"class":"ap-answer-content ap-q-content"})
            for x in table:
                temp = x.findAll('p')
                for y in temp:
                    t = y.text
                    p2.add_run(t+"\n")

    document2.save("giai.docx")

get("http://hoctap.dvtienich.com/2019-thi-thu-thpt-quoc-gia-truong-do-luong-1-nghe-an-lan-3/")